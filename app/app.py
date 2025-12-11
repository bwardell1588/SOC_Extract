import os
import io
import json
import uuid
import time
import hashlib
import logging
import re
from json import JSONDecodeError
from dataclasses import dataclass
from typing import Dict, Any, List, Tuple

from flask import Flask, request, jsonify, send_file, render_template_string

import boto3
from botocore.config import Config
from botocore.exceptions import BotoCoreError, EndpointConnectionError, ReadTimeoutError

from pdfminer.high_level import extract_text_to_fp
from pdfminer.layout import LAParams

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.section import WD_ORIENT
except Exception:
    Document = None
    Inches = Pt = WD_ORIENT = None

# -----------------------------
# Logging
# -----------------------------

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger("soc2")

# -----------------------------
# App + Config
# -----------------------------

app = Flask(__name__)

DEFAULT_REPORT_DIR = r"C:\Users\brian.wardell\AI_Projects\soc_extract\app\generated_reports"


@dataclass
class AppConfig:
    BEDROCK_REGION: str = os.getenv("BEDROCK_REGION", "us-east-1")
    BEDROCK_MODEL_ID: str = os.getenv(
        "BEDROCK_MODEL_ID",
        "global.anthropic.claude-haiku-4-5-20251001-v1:0"
    )
    READ_TIMEOUT: int = int(os.getenv("READ_TIMEOUT", "120"))
    CONNECT_TIMEOUT: int = int(os.getenv("CONNECT_TIMEOUT", "15"))
    MAX_TOTAL_TOKENS: int = int(os.getenv("MAX_TOTAL_TOKENS", "10000"))
    TEMPERATURE: float = float(os.getenv("TEMPERATURE", "0.0"))
    REPORT_DIR: str = os.getenv("REPORT_DIR", DEFAULT_REPORT_DIR)
    BATCH_SIZE: int = int(os.getenv("BATCH_SIZE", "20"))
    MAX_RETRIES: int = int(os.getenv("MAX_RETRIES", "3"))
    LOG_BEDROCK_FULL: bool = os.getenv("LOG_BEDROCK_FULL", "1") not in ("0", "false", "False", "")
    LOG_BEDROCK_MAX_CHARS: int = int(os.getenv("LOG_BEDROCK_MAX_CHARS", "0"))


app.config.from_object(AppConfig)

# -----------------------------
# Reports dir helpers
# -----------------------------


def _resolve_report_dir() -> str:
    base = (app.config.get("REPORT_DIR") or DEFAULT_REPORT_DIR).strip()
    try:
        os.makedirs(base, exist_ok=True)
    except Exception as e:
        import tempfile
        fallback = os.path.join(tempfile.gettempdir(), "soc2_reports")
        os.makedirs(fallback, exist_ok=True)
        logger.warning("Could not create REPORT_DIR %r (%s). Falling back to %r", base, e, fallback)
        base = fallback
    return base


def _report_path_docx(report_id: str) -> str:
    return os.path.join(_resolve_report_dir(), f"soc2_report_{report_id}.docx")


# Cache: doc_id -> {full_text, table_text, narrative_text}
PARSED_CACHE: Dict[str, Dict[str, Any]] = {}

# -----------------------------
# AWS Credentials
# -----------------------------


def _env_creds_from_custom_vars():
    ak = os.getenv("AWS_ACCESS_KEY")
    sk = os.getenv("AWS_SECRET")
    if ak and sk:
        return {"aws_access_key_id": ak, "aws_secret_access_key": sk}
    return {}


# -----------------------------
# Bedrock client
# -----------------------------


def _new_bedrock_client():
    cfg = Config(
        region_name=app.config["BEDROCK_REGION"],
        retries={"max_attempts": 3, "mode": "standard"},
        read_timeout=app.config["READ_TIMEOUT"],
        connect_timeout=app.config["CONNECT_TIMEOUT"],
    )
    return boto3.client("bedrock-runtime", config=cfg, **_env_creds_from_custom_vars())


# -----------------------------
# JSON parsing helpers
# -----------------------------


def _strip_code_fences(text: str) -> str:
    if not text:
        return text
    s = text.strip()
    s = re.sub(r"^```(?:json)?\s*", "", s, flags=re.IGNORECASE)
    s = re.sub(r"\s*```$", "", s)
    return s.strip()


def _parse_model_json(text: str):
    try:
        return json.loads(text)
    except JSONDecodeError:
        pass
    s = _strip_code_fences(text)
    try:
        return json.loads(s)
    except JSONDecodeError:
        pass
    start, end = s.find("{"), s.rfind("}")
    if start != -1 and end != -1 and end > start:
        return json.loads(s[start:end + 1])
    raise JSONDecodeError("Model text was not valid JSON", doc=s, pos=0)


# -----------------------------
# Bedrock invocation
# -----------------------------


def _maybe_truncate(s: str) -> str:
    max_chars = app.config["LOG_BEDROCK_MAX_CHARS"]
    if max_chars and len(s) > max_chars:
        return s[:max_chars] + f"... [truncated to {max_chars} chars]"
    return s


def invoke_bedrock(system_text: str, instruction_text: str, use_cache: bool = True) -> Dict[str, Any]:
    client = _new_bedrock_client()
    
    # Build system content - with or without cache_control
    if use_cache:
        system_content = [{
            "type": "text",
            "text": system_text,
            "cache_control": {"type": "ephemeral"}
        }]
    else:
        system_content = [{
            "type": "text",
            "text": system_text
        }]
    
    request_body = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": app.config["MAX_TOTAL_TOKENS"],
        "temperature": app.config["TEMPERATURE"],
        "system": system_content,
        "messages": [{"role": "user", "content": [{"type": "text", "text": instruction_text}]}]
    }

    resp = client.invoke_model(
        modelId=app.config["BEDROCK_MODEL_ID"],
        body=json.dumps(request_body)
    )

    raw_body = resp.get("body")
    if not raw_body:
        raise RuntimeError("Bedrock returned empty body.")

    try:
        body_json = json.loads(raw_body.read())
    except Exception as e:
        raise RuntimeError(f"Could not decode Bedrock response: {e}")

    if app.config["LOG_BEDROCK_FULL"]:
        logger.info("[Bedrock] Response:\n%s", _maybe_truncate(json.dumps(body_json, indent=2)))

    contents = body_json.get("content", []) if isinstance(body_json, dict) else []
    text_blocks = [c.get("text") for c in contents if isinstance(c, dict) and c.get("type") == "text"]
    model_text = "".join(filter(None, text_blocks)).strip()

    if not model_text:
        raise RuntimeError("No text content in Bedrock response")

    return _parse_model_json(model_text)


def invoke_bedrock_with_retry(system_text: str, instruction_text: str, use_cache: bool = True) -> Dict[str, Any]:
    last_err = None
    for attempt in range(app.config["MAX_RETRIES"]):
        try:
            return invoke_bedrock(system_text, instruction_text, use_cache=use_cache)
        except (ReadTimeoutError, EndpointConnectionError, BotoCoreError, RuntimeError, JSONDecodeError) as e:
            last_err = e
            logger.warning("Bedrock attempt %d failed: %s", attempt + 1, e)
            time.sleep(2 ** attempt)
    raise RuntimeError(f"Bedrock failed after {app.config['MAX_RETRIES']} attempts: {last_err}")


# -----------------------------
# PDF extraction and segmentation
# -----------------------------


def extract_pdf_text(file_bytes: bytes) -> Tuple[str, List[str]]:
    """Extract text from PDF, return (full_labeled_text, list_of_page_texts)."""
    output = io.StringIO()
    extract_text_to_fp(io.BytesIO(file_bytes), output, laparams=LAParams(), output_type="text")
    raw = output.getvalue()

    # Split on form feeds (page breaks)
    pages = raw.split("\f") if "\f" in raw else [raw]
    pages = [p.strip() for p in pages]

    # Create labeled full text with simplified markers
    labeled = [f"=== PAGE {i} ===\n{p}" for i, p in enumerate(pages, 1)]
    full_text = "\n\n".join(labeled)

    return full_text, pages


def detect_section_markers(pages: List[str]) -> Dict[int, int]:
    """
    Detect SOC report section boundaries.
    Returns a dict mapping page indices to section numbers.
    
    Section detection logic:
    - Looks for "SECTION I", "SECTION II", "SECTION III", etc. patterns
    - Handles multiple formats:
      - "SECTION III" standalone on its own line
      - "SECTION III - Title" with dash separator
      - "SECTION 3" numeric formats
      - "SECTION THREE" spelled-out numbers
      - "SECTION - THREE" with dash before number
      - "III." standalone Roman numeral with period
    - Forward-fills section numbers to subsequent pages
    - Ignores section references in Table of Contents
    """
    section_info: Dict[int, int] = {}
    current_section = 0
    
    # Roman numeral mapping
    roman_to_int = {
        'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5, 
        'VI': 6, 'VII': 7, 'VIII': 8, 'IX': 9, 'X': 10
    }
    
    # Spelled-out number mapping
    word_to_int = {
        'ONE': 1, 'TWO': 2, 'THREE': 3, 'FOUR': 4, 'FIVE': 5,
        'SIX': 6, 'SEVEN': 7, 'EIGHT': 8, 'NINE': 9, 'TEN': 10
    }
    
    def parse_section_number(section_str):
        """Convert section string (roman, arabic, or word) to integer."""
        section_str = section_str.upper().strip()
        if section_str in roman_to_int:
            return roman_to_int[section_str]
        if section_str in word_to_int:
            return word_to_int[section_str]
        if section_str.isdigit():
            return int(section_str)
        return 0
    
    # Combined pattern for section numbers: roman numerals, digits, or spelled-out words
    section_num_pattern = r'([IVX]+|[0-9]+|ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN)'
    
    for i, page in enumerate(pages):
        detected_section = 0
        
        # Skip Table of Contents pages - they mention all sections but aren't the actual sections
        if 'Table of Contents' in page[:500] or 'TABLE OF CONTENTS' in page[:500]:
            section_info[i] = current_section
            continue
        
        # Pattern 1: "SECTION III" or "SECTION THREE" on its own line (standalone header)
        standalone_pattern = re.compile(
            r'^\s*SECTION\s+' + section_num_pattern + r'\s*$', 
            re.MULTILINE | re.IGNORECASE
        )
        match = standalone_pattern.search(page)
        if match:
            detected_section = parse_section_number(match.group(1))
        
        # Pattern 2: "SECTION III - Title" or "SECTION THREE - Title" format (with dash and title)
        if not detected_section:
            dash_after_pattern = re.compile(
                r'^\s*SECTION\s+' + section_num_pattern + r'\s*[-–—]\s*\S', 
                re.MULTILINE | re.IGNORECASE
            )
            match = dash_after_pattern.search(page)
            if match:
                detected_section = parse_section_number(match.group(1))
        
        # Pattern 3: "SECTION - THREE" or "SECTION - III" format (with dash before number)
        if not detected_section:
            dash_before_pattern = re.compile(
                r'^\s*SECTION\s*[-–—]\s*' + section_num_pattern + r'(?:\s|$|[-–—])', 
                re.MULTILINE | re.IGNORECASE
            )
            match = dash_before_pattern.search(page)
            if match:
                detected_section = parse_section_number(match.group(1))
        
        # Pattern 4: "SECTION III" or "SECTION THREE" followed by newline (in running text)
        if not detected_section:
            inline_pattern = re.compile(
                r'SECTION\s+' + section_num_pattern + r'\s*[\n\r]', 
                re.IGNORECASE
            )
            match = inline_pattern.search(page)
            if match:
                detected_section = parse_section_number(match.group(1))
        
        # Pattern 5: Standalone "III." or "IV." (Roman numeral with period, no "SECTION" word)
        # Must be at start of line to avoid matching list items
        if not detected_section:
            roman_standalone_pattern = re.compile(
                r'^\s*([IVX]+)\.\s*$',
                re.MULTILINE
            )
            match = roman_standalone_pattern.search(page)
            if match:
                detected_section = parse_section_number(match.group(1))
        
        # Pattern 6: "III. Title" or "IV. Title" (Roman numeral with period followed by title)
        if not detected_section:
            roman_title_pattern = re.compile(
                r'^\s*([IVX]+)\.\s+[A-Z]',
                re.MULTILINE
            )
            match = roman_title_pattern.search(page)
            if match:
                # Only match if it looks like a section header (near top of page or after whitespace)
                match_pos = match.start()
                # Check if this is near the beginning of meaningful content
                preceding_text = page[:match_pos].strip()
                # Accept if it's near the top or after minimal content (like headers/footers)
                if len(preceding_text) < 200:
                    detected_section = parse_section_number(match.group(1))
        
        # Update current section if we found a new one
        if detected_section > 0:
            current_section = detected_section
        
        section_info[i] = current_section
    
    return section_info


def analyze_text_structure(page_text: str) -> Dict[str, Any]:
    """
    Analyze structural characteristics of page text.
    Used as fallback scoring when section markers aren't available.
    """
    lines = [ln for ln in page_text.split("\n") if ln.strip()]
    text_lower = page_text.lower()
    
    analysis: Dict[str, Any] = {
        'line_count': len(lines),
        'indicators': [],
        'score_adjustments': 0
    }
    
    # Test result language density (strong indicator of control testing tables)
    result_phrases = [
        "no exceptions noted",
        "inquired of",
        "inspected the",
        "determined that",
        "observed that"
    ]
    result_density = sum(text_lower.count(phrase) for phrase in result_phrases)
    if result_density >= 3:
        analysis['score_adjustments'] += 2
        analysis['indicators'].append(f"test result density: {result_density}")
    
    # CC criteria references (e.g., CC1.1, CC6.2)
    criteria_pattern = re.compile(r'\bCC\s*\d+\.\d+\b', re.IGNORECASE)
    criteria_matches = criteria_pattern.findall(page_text)
    if len(criteria_matches) >= 2:
        analysis['score_adjustments'] += 1
        analysis['indicators'].append(f"CC criteria: {len(criteria_matches)}")
    
    # Table continuation markers
    if "(continued)" in text_lower:
        analysis['score_adjustments'] += 1
        analysis['indicators'].append("continued marker")
    
    # Multi-column layout detection
    multi_column_lines = 0
    for line in lines:
        chunks = re.split(r'\s{3,}', line.strip())
        chunks = [c for c in chunks if c.strip()]
        if len(chunks) >= 2:
            multi_column_lines += 1
    
    if lines and multi_column_lines / len(lines) > 0.3:
        analysis['score_adjustments'] += 1
        analysis['indicators'].append(f"multi-column lines: {multi_column_lines}")
    
    return analysis


def is_table_page(page_text: str, page_idx: int, section_info: Dict[int, int]) -> Dict[str, Any]:
    """
    Determine if a page should be classified as table content.
    
    Primary logic (section-based):
    - Sections I-II (auditor opinion, management assertion) -> NARRATIVE
    - Section 0 (before any section marker, e.g., cover, TOC) -> NARRATIVE
    - Sections III+ (system description, controls, testing) -> TABLE
    
    Fallback logic (when section not detected but after Section III has started):
    - Use structural analysis
    """
    reasons: List[str] = []
    score = 0
    
    section_num = section_info.get(page_idx, 0)
    
    # ===========================================
    # PRIMARY RULE: Section-based classification
    # ===========================================
    
    # Section 0, I, and II are ALWAYS narrative
    # Section 0 = pages before any section marker (cover page, TOC, etc.)
    if section_num in [0, 1, 2]:
        if section_num == 0:
            reasons.append("Before Section III (cover/TOC -> narrative)")
        else:
            reasons.append(f"Section {section_num} (auditor/management -> narrative)")
        return {
            "is_table": False,
            "score": -10,
            "reasons": reasons,
            "forced_narrative": True,
            "forced_table": False
        }
    
    # Sections III+ are ALWAYS table content
    if section_num >= 3:
        reasons.append(f"Section {section_num} (system/controls/testing -> table)")
        return {
            "is_table": True,
            "score": 10,
            "reasons": reasons,
            "forced_narrative": False,
            "forced_table": True
        }
    
    # ===========================================
    # FALLBACK: Should rarely reach here
    # Use structural analysis
    # ===========================================
    reasons.append("Unknown section - using structural analysis")
    
    # Text structure analysis
    text_analysis = analyze_text_structure(page_text)
    score += text_analysis['score_adjustments']
    reasons.extend(text_analysis['indicators'])
    
    # Threshold for table classification
    is_table_result = score >= 2
    
    return {
        "is_table": is_table_result,
        "score": score,
        "reasons": reasons,
        "forced_narrative": False,
        "forced_table": False
    }


def segment_content(pages: List[str]) -> Dict[str, Any]:
    """
    Separate table content from narrative content using section-based detection.
    
    Classification rules:
    - Sections I-II (auditor opinion, management assertion) -> NARRATIVE
    - Sections III+ (system description, controls, testing) -> TABLE
    """
    
    # Detect section boundaries
    section_info = detect_section_markers(pages)
    
    page_classifications: List[Dict[str, Any]] = []

    # Classify each page based on section
    for i, page in enumerate(pages):
        analysis = is_table_page(page, i, section_info)
        page_classifications.append({
            "page": i + 1,
            "is_table": analysis["is_table"],
            "section": section_info.get(i, 0),
            "forced_narrative": analysis.get("forced_narrative", False),
            "forced_table": analysis.get("forced_table", False),
            "reasons": analysis["reasons"],
        })
    
    # Build the text segments
    table_parts: List[str] = []
    narrative_parts: List[str] = []

    for i, classification in enumerate(page_classifications):
        labeled = f"=== PAGE {classification['page']} ===\n{pages[i]}"
        if classification["is_table"]:
            table_parts.append(labeled)
        else:
            narrative_parts.append(labeled)

    table_text = "\n\n".join(table_parts)
    narrative_text = "\n\n".join(narrative_parts)
    
    # Find where Section III starts for logging
    sec3_start = next((c['page'] for c in page_classifications if c['section'] >= 3), None)

    logger.info("[Segment] Section-based: %d table pages (Section III+), %d narrative pages (Sections 0-II). Section III starts at page %s", 
                len(table_parts), len(narrative_parts), sec3_start)

    return {
        "table_text": table_text,
        "narrative_text": narrative_text,
        "classifications": page_classifications,
    }



# -----------------------------
# Control merging/deduplication
# -----------------------------


def _control_key(c: Dict[str, Any]) -> str:
    title = (c.get("control_title") or "").strip().lower()
    desc = (c.get("control_description") or "").strip().lower()
    return hashlib.sha1(f"{title}|||{desc}".encode()).hexdigest()


def merge_controls(existing: List[Dict[str, Any]], new_ones: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    seen = {_control_key(x) for x in existing}
    result = list(existing)

    for c in new_ones or []:
        key = _control_key(c)
        if key not in seen:
            seen.add(key)
            result.append(c)

    # Assign IDs to controls missing them
    next_id = 1
    for c in result:
        if not (c.get("control_id") or "").strip():
            c["control_id"] = f"C-{next_id:03d}"
            next_id += 1

    return result


def merge_criteria_into_controls(
    controls: List[Dict[str, Any]], 
    criteria_mappings: List[Dict[str, Any]]
) -> List[Dict[str, Any]]:
    """
    Merge criteria mappings into controls.
    Creates a reverse lookup: control_id -> list of criteria that reference it.
    
    IMPORTANT: Preserves any criteria already extracted with the control,
    only adds new criteria from the mapping that aren't already present.
    """
    # Build reverse mapping: control_id -> [criterion_ids]
    control_to_criteria: Dict[str, List[str]] = {}
    
    for mapping in criteria_mappings:
        criterion_id = mapping.get("criterion_id", "")
        mapped_controls = mapping.get("mapped_controls", [])
        
        for control_ref in mapped_controls:
            # Normalize control reference (strip whitespace)
            control_ref = str(control_ref).strip()
            if control_ref:
                if control_ref not in control_to_criteria:
                    control_to_criteria[control_ref] = []
                if criterion_id not in control_to_criteria[control_ref]:
                    control_to_criteria[control_ref].append(criterion_id)
    
    logger.info("[Merge] Built reverse mapping for %d control references", len(control_to_criteria))
    
    # Track stats
    controls_with_existing_criteria = 0
    controls_with_added_criteria = 0
    
    # Apply criteria to each control (merge, don't replace)
    for control in controls:
        control_id = (control.get("control_id") or "").strip()
        
        # Get existing criteria from control extraction (may be empty list)
        existing_criteria = control.get("criterion") or []
        if isinstance(existing_criteria, str):
            existing_criteria = [existing_criteria] if existing_criteria else []
        existing_criteria = [c.strip() for c in existing_criteria if c.strip()]
        
        if existing_criteria:
            controls_with_existing_criteria += 1
        
        # Find additional criteria from mapping tables
        mapped_criteria = control_to_criteria.get(control_id, [])
        
        # Also try normalized matching if no exact match
        if not mapped_criteria:
            for ref in control_to_criteria:
                if _normalize_control_ref(ref) == _normalize_control_ref(control_id):
                    mapped_criteria = control_to_criteria[ref]
                    break
        
        # Merge: start with existing, add any from mapping that aren't already present
        existing_set = set(existing_criteria)
        new_criteria = [c for c in mapped_criteria if c not in existing_set]
        
        if new_criteria:
            controls_with_added_criteria += 1
        
        # Combined and sorted
        all_criteria = existing_criteria + new_criteria
        control["criterion"] = sorted(set(all_criteria)) if all_criteria else []
    
    logger.info("[Merge] %d controls had criteria from extraction, %d had criteria added from mapping",
                controls_with_existing_criteria, controls_with_added_criteria)
    
    return controls


def _normalize_control_ref(ref: str) -> str:
    """Normalize control reference for matching (e.g., '1.01' -> '1.1')."""
    ref = str(ref).strip()
    # Try to parse as X.Y and normalize
    match = re.match(r'^(\d+)\.(\d+)$', ref)
    if match:
        major, minor = match.groups()
        return f"{int(major)}.{int(minor)}"
    return ref


# -----------------------------
# DOCX report builder
# -----------------------------


def _apply_table_font(table, font_size_pt: int = 8) -> None:
    """Apply consistent font size to all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size_pt)


def _set_cell_border(cell, border_type: str = "none"):
    """Set cell borders. Use 'none' to hide borders."""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    
    if border_type == "none":
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            r'<w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/>'
            r'</w:tcBorders>'
        )
        tcPr.append(tcBorders)


def build_report_docx(
    auditor_opinion: Dict[str, Any],
    subservice_controls: List[Dict[str, Any]],
    user_entity_controls: List[Dict[str, Any]],
    vendor_controls: List[Dict[str, Any]],
    exceptions: List[Dict[str, Any]],
    criteria_mappings: List[Dict[str, Any]],
    out_path: str
) -> None:
    if Document is None:
        raise RuntimeError("python-docx required: pip install python-docx")

    logger.info("[Report] Building DOCX: %s (vendor=%d, subservice=%d, user_entity=%d, exceptions=%d, criteria=%d)", 
                out_path, len(vendor_controls or []), len(subservice_controls or []), 
                len(user_entity_controls or []), len(exceptions or []), len(criteria_mappings or []))
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)

    doc = Document()

    # Landscape with 0.5" margins
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = section.bottom_margin = Inches(0.5)
    section.left_margin = section.right_margin = Inches(0.5)

    doc.add_heading("SOC 2 Report Extraction", 0)

    # ===========================================
    # Section 1: General Info Table (no borders)
    # ===========================================
    doc.add_heading("General Information", level=1)

    info_table = doc.add_table(rows=6, cols=2)
    
    # Row 0: Service/Product
    info_table.rows[0].cells[0].text = "Service/Product:"
    info_table.rows[0].cells[1].text = auditor_opinion.get("service_product") or ""
    
    # Row 1: Report Type
    info_table.rows[1].cells[0].text = "Report Type:"
    info_table.rows[1].cells[1].text = auditor_opinion.get("report_type") or ""
    
    # Row 2: Scope Date
    info_table.rows[2].cells[0].text = "Scope Date:"
    info_table.rows[2].cells[1].text = auditor_opinion.get("scope_date") or ""
    
    # Row 3: Auditor's Name
    info_table.rows[3].cells[0].text = "Auditor's Name:"
    info_table.rows[3].cells[1].text = auditor_opinion.get("auditors_name") or ""
    
    # Row 4: Qualified Opinion
    info_table.rows[4].cells[0].text = "Qualified Opinion:"
    qualified = auditor_opinion.get("qualified_opinion")
    info_table.rows[4].cells[1].text = "Yes" if qualified else "No" if qualified is False else ""
    
    # Row 5: Auditor's Opinion
    info_table.rows[5].cells[0].text = "Auditor's Opinion:"
    info_table.rows[5].cells[1].text = auditor_opinion.get("auditors_opinion") or ""
    
    # Remove borders from info table
    for row in info_table.rows:
        for cell in row.cells:
            _set_cell_border(cell, "none")
    
    # Set column widths
    for row in info_table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(8.5)
    
    _apply_table_font(info_table, font_size_pt=10)

    # ===========================================
    # Section 2: Criteria Mappings
    # ===========================================
    doc.add_paragraph()
    doc.add_heading("Criteria Mappings", level=1)

    crit_headers = ["Criterion ID", "Description", "Mapped Controls"]
    crit_col_widths = [Inches(1.2), Inches(5.0), Inches(3.8)]

    crit_table = doc.add_table(rows=1, cols=len(crit_headers))
    crit_hdr = crit_table.rows[0].cells
    for i, h in enumerate(crit_headers):
        crit_hdr[i].text = h
        crit_hdr[i].width = crit_col_widths[i]

    if criteria_mappings:
        for mapping in criteria_mappings:
            row = crit_table.add_row().cells
            row[0].text = mapping.get("criterion_id") or ""
            row[1].text = mapping.get("criterion_description") or ""
            mapped = mapping.get("mapped_controls") or []
            row[2].text = ", ".join(mapped) if isinstance(mapped, list) else str(mapped)
            for i, cell in enumerate(row):
                cell.width = crit_col_widths[i]
    else:
        row = crit_table.add_row().cells
        row[0].text = "No criteria mappings found"

    _apply_table_font(crit_table)

    # ===========================================
    # Section 3: Complementary Subservice Organization Controls
    # ===========================================
    doc.add_paragraph()
    doc.add_heading("Complementary Subservice Organization Controls", level=1)

    sub_headers = ["Organization", "Control ID", "Description", "Criteria Covered"]
    sub_col_widths = [Inches(1.5), Inches(0.8), Inches(5.0), Inches(2.5)]

    sub_table = doc.add_table(rows=1, cols=len(sub_headers))
    sub_hdr = sub_table.rows[0].cells
    for i, h in enumerate(sub_headers):
        sub_hdr[i].text = h
        sub_hdr[i].width = sub_col_widths[i]

    has_subservice_rows = False
    for c in subservice_controls or []:
        has_subservice_rows = True
        row = sub_table.add_row().cells
        row[0].text = c.get("organization_name") or c.get("name") or ""
        row[1].text = c.get("control_id") or ""
        row[2].text = c.get("description") or ""
        criteria = c.get("criteria_covered") or []
        row[3].text = ", ".join(criteria) if isinstance(criteria, list) else str(criteria)
        for i, cell in enumerate(row):
            cell.width = sub_col_widths[i]

    if not has_subservice_rows:
        row = sub_table.add_row().cells
        row[0].text = "No subservice organization controls found"

    _apply_table_font(sub_table)

    # ===========================================
    # Section 4: Complementary User Entity Controls
    # ===========================================
    doc.add_paragraph()
    doc.add_heading("Complementary User Entity Controls", level=1)

    ue_headers = ["Category", "Control ID", "Description", "Criteria Covered"]
    ue_col_widths = [Inches(1.5), Inches(0.8), Inches(5.0), Inches(2.5)]

    ue_table = doc.add_table(rows=1, cols=len(ue_headers))
    ue_hdr = ue_table.rows[0].cells
    for i, h in enumerate(ue_headers):
        ue_hdr[i].text = h
        ue_hdr[i].width = ue_col_widths[i]

    has_user_entity_rows = False
    for c in user_entity_controls or []:
        has_user_entity_rows = True
        row = ue_table.add_row().cells
        row[0].text = c.get("category") or c.get("name") or ""
        row[1].text = c.get("control_id") or ""
        row[2].text = c.get("description") or ""
        criteria = c.get("criteria_covered") or []
        row[3].text = ", ".join(criteria) if isinstance(criteria, list) else str(criteria)
        for i, cell in enumerate(row):
            cell.width = ue_col_widths[i]

    if not has_user_entity_rows:
        row = ue_table.add_row().cells
        row[0].text = "No user entity controls found"

    _apply_table_font(ue_table)

    # ===========================================
    # Section 5: Vendor/Service Organization Controls
    # ===========================================
    doc.add_paragraph()
    doc.add_heading("Vendor/Service Organization Controls", level=1)

    headers = ["Control ID", "Criterion", "Title", "Description", "Tests Applied", "Result"]
    col_widths = [Inches(0.8), Inches(1.2), Inches(1.8), Inches(2.8), Inches(2.8), Inches(0.8)]

    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].width = col_widths[i]

    for c in vendor_controls or []:
        row = table.add_row().cells
        row[0].text = c.get("control_id") or ""
        criterion = c.get("criterion") or []
        row[1].text = ", ".join(criterion) if isinstance(criterion, list) else str(criterion)
        row[2].text = c.get("control_title") or ""
        row[3].text = c.get("control_description") or ""
        row[4].text = ", ".join(c.get("tests_applied") or [])
        row[5].text = c.get("result") or ""
        for i, cell in enumerate(row):
            cell.width = col_widths[i]

    if not vendor_controls:
        row = table.add_row().cells
        row[2].text = "No vendor controls extracted"

    _apply_table_font(table)

    # ===========================================
    # Section 5: Exceptions
    # ===========================================
    doc.add_paragraph()
    doc.add_heading("Exceptions", level=1)

    exc_headers = ["Control Objective", "Testing Description", "Exception Description", "Management Response"]
    exc_col_widths = [Inches(1.5), Inches(2.5), Inches(3.0), Inches(3.0)]

    exc_table = doc.add_table(rows=1, cols=len(exc_headers))
    exc_hdr = exc_table.rows[0].cells
    for i, h in enumerate(exc_headers):
        exc_hdr[i].text = h
        exc_hdr[i].width = exc_col_widths[i]

    if exceptions:
        for e in exceptions:
            row = exc_table.add_row().cells
            row[0].text = e.get("control_objective") or ""
            row[1].text = e.get("testing_description") or ""
            row[2].text = e.get("exception_description") or ""
            row[3].text = e.get("management_response") or ""
            for i, cell in enumerate(row):
                cell.width = exc_col_widths[i]
    else:
        row = exc_table.add_row().cells
        row[0].text = "No exceptions found"

    _apply_table_font(exc_table)

    doc.save(out_path)
    logger.info("[Report] Wrote: %s (%d bytes)", out_path, os.path.getsize(out_path))


# -----------------------------
# HTML UI
# -----------------------------

INDEX_HTML = r'''
<!doctype html>
<html>
  <head>
    <title>SOC 2 Control Extractor</title>
    <style>
      body { font-family: system-ui, -apple-system, sans-serif; margin: 40px; }
      .card { border: 1px solid #ddd; border-radius: 10px; padding: 20px; max-width: 900px; }
      .row { margin-top: 12px; }
      .muted { color: #666; }
      button { padding: 10px 16px; border-radius: 8px; border: 1px solid #444; background: #111; color: #fff; cursor: pointer; }
      button:disabled { background: #aaa; cursor: not-allowed; }
      pre { background: #f6f8fa; padding: 12px; border-radius: 6px; overflow: auto; }
      .badges { display: flex; gap: 8px; margin: 10px 0; }
      .badge { padding: 6px 10px; border-radius: 999px; border: 1px solid #aaa; font-size: 12px; }
      .off { background: #f3f4f6; color: #6b7280; }
      .on { background: #e0f2fe; color: #0369a1; }
      .ok { background: #dcfce7; color: #166534; }
      .err { background: #fee2e2; color: #991b1b; }
      a.download { display: inline-block; margin-top: 10px; margin-right: 12px; }
    </style>
  </head>
  <body>
    <div class="card">
      <h1>SOC 2 Control Extractor</h1>
      <p class="muted">Upload a SOC 2 PDF to extract controls and download a DOCX report.</p>

      <div class="row">
        <form id="uploadForm">
          <input type="file" id="pdf" name="pdf" accept="application/pdf" required />
          <button id="submitBtn" type="submit">Run Extraction</button>
        </form>
      </div>

      <div class="row badges">
        <div id="badgeParsed" class="badge off">Parsed</div>
        <div id="badgeExtracting" class="badge off">Extracting…</div>
        <div id="badgeDone" class="badge off">Done</div>
      </div>

      <div class="row" id="statusText"></div>
      <div class="row" id="result"></div>
    </div>

    <script>
      const form = document.getElementById('uploadForm');
      const fileInput = document.getElementById('pdf');
      const resultDiv = document.getElementById('result');
      const btn = document.getElementById('submitBtn');
      const statusText = document.getElementById('statusText');
      const badges = {
        parsed: document.getElementById('badgeParsed'),
        extracting: document.getElementById('badgeExtracting'),
        done: document.getElementById('badgeDone')
      };

      function setBadge(el, state) {
        el.classList.remove('off', 'on', 'ok', 'err');
        el.classList.add(state);
      }

      function resetUI() {
        resultDiv.innerHTML = '';
        statusText.textContent = '';
        Object.values(badges).forEach(b => setBadge(b, 'off'));
      }

      fileInput.addEventListener('change', resetUI);

      form.addEventListener('submit', async (e) => {
        e.preventDefault();
        resetUI();
        btn.disabled = true;

        try {
          const formData = new FormData();
          formData.append('pdf', fileInput.files[0]);

          const parseResp = await fetch('/parse', { method: 'POST', body: formData });
          const parseData = await parseResp.json();
          if (!parseData.ok) {
            setBadge(badges.parsed, 'err');
            resultDiv.innerHTML = '<p style="color:red;">' + (parseData.error || 'Parse failed') + '</p>';
            return;
          }
          setBadge(badges.parsed, 'ok');

          setBadge(badges.extracting, 'on');
          statusText.textContent = 'Extracting controls...';

          const extractResp = await fetch('/extract_cursor/' + parseData.doc_id, { method: 'POST' });
          const extractData = await extractResp.json();
          if (!extractData.ok) {
            setBadge(badges.extracting, 'err');
            resultDiv.innerHTML = '<p style="color:red;">' + (extractData.error || 'Extraction failed') + '</p>';
            return;
          }

          setBadge(badges.extracting, 'off');
          setBadge(badges.done, 'ok');
          statusText.textContent = '';

          const pre = document.createElement('pre');
          pre.textContent = JSON.stringify(extractData.result, null, 2);
          resultDiv.appendChild(pre);

          if (extractData.report_id) {
            const link = document.createElement('a');
            link.href = '/download/docx/' + extractData.report_id;
            link.textContent = 'Download DOCX Report';
            link.className = 'download';
            resultDiv.appendChild(link);
          }
        } catch (err) {
          setBadge(badges.extracting, 'err');
          resultDiv.innerHTML = '<p style="color:red;">' + (err?.message || err) + '</p>';
        } finally {
          btn.disabled = false;
        }
      });
    </script>
  </body>
</html>
'''


@app.get("/")
def index():
    return render_template_string(INDEX_HTML)


# -----------------------------
# Routes
# -----------------------------

from prompts import (
    SYSTEM_DOCUMENT_TABLES,
    SYSTEM_DOCUMENT_GENERAL,
    INSTRUCTION_CONTROLS,
    INSTRUCTION_SUBSERVICE,
    INSTRUCTION_USER_ENTITY,
    INSTRUCTION_CRITERIA,
    INSTRUCTION_EXCEPTIONS,
    INSTRUCTION_AUDITOR_OPINION,
)


@app.post("/parse")
def parse_pdf():
    """Parse PDF and segment into table/narrative content."""
    PARSED_CACHE.clear()

    if "pdf" not in request.files:
        return jsonify({"ok": False, "error": "No PDF uploaded"}), 400

    pdf_file = request.files["pdf"]
    raw = pdf_file.read()
    if not raw:
        return jsonify({"ok": False, "error": "Could not read PDF"}), 400

    try:
        full_text, pages = extract_pdf_text(raw)
        segments = segment_content(pages)
    except Exception as e:
        logger.exception("PDF parse error")
        return jsonify({"ok": False, "error": f"PDF parse error: {e}"}), 500

    doc_id = str(uuid.uuid4())[:8]
    PARSED_CACHE[doc_id] = {
        "full_text": full_text,
        "table_text": segments["table_text"],
        "narrative_text": segments["narrative_text"],
        "pages": pages,  # Store individual pages for partial extraction
    }

    # Log stats
    table_pages = segments["table_text"].count("=== PAGE")
    narrative_pages = segments["narrative_text"].count("=== PAGE")
    logger.info(
        "[Parse] doc_id=%s, total_pages=%d, table_pages=%d, narrative_pages=%d, full_chars=%d, table_chars=%d",
        doc_id, len(pages), table_pages, narrative_pages, len(full_text), len(segments["table_text"])
    )

    return jsonify({
        "ok": True,
        "doc_id": doc_id,
        "total_pages": len(pages),
        "table_pages": table_pages,
        "narrative_pages": narrative_pages,
        "full_chars": len(full_text),
        "table_chars": len(segments["table_text"]),
    })


@app.post("/extract_cursor/<doc_id>")
def extract_controls_cursor(doc_id: str):
    """
    Multi-phase extraction with shared system prompts for cache reuse:
    1. Extract auditor's opinion (from first 25 pages, no caching)
    2. Extract vendor/service organization controls (batched, cursor-based, cached)
    3. Extract exceptions (single call, cached)
    4. Extract subservice organization controls (single call, cached)
    5. Extract user entity controls (single call, cached)
    6. Extract criteria mappings (single call, cached)
    7. Merge criteria into vendor controls
    8. Generate DOCX with all content
    """
    if doc_id not in PARSED_CACHE:
        return jsonify({"ok": False, "error": "Unknown doc_id"}), 404

    cache = PARSED_CACHE[doc_id]

    # Get text segments
    pages = cache.get("pages", [])
    table_text = cache["table_text"] if cache["table_text"].strip() else cache["full_text"]
    
    # Build first 25 pages text for auditor opinion (no caching needed)
    first_pages = pages[:25]
    first_pages_text = "\n".join([f"=== PAGE {i+1} ===\n{p}" for i, p in enumerate(first_pages)])
    
    logger.info("[Extract] Using first %d pages (%d chars) for auditor opinion, table_text (%d chars) for controls", 
                len(first_pages), len(first_pages_text), len(table_text))

    t0 = time.time()

    # System prompt for tables (cached for multiple calls)
    system_prompt_tables = SYSTEM_DOCUMENT_TABLES.replace("{{DOCUMENT_TEXT}}", table_text)
    
    # System prompt for auditor opinion (first 25 pages, not cached)
    system_prompt_opinion = SYSTEM_DOCUMENT_GENERAL.replace("{{DOCUMENT_TEXT}}", first_pages_text)

    # ===========================================
    # Phase 1: Extract auditor's opinion (first 25 pages, no cache)
    # ===========================================
    logger.info("[Extract Auditor Opinion] Starting (first 25 pages, no cache)...")
    
    try:
        opinion_result = invoke_bedrock_with_retry(system_prompt_opinion, INSTRUCTION_AUDITOR_OPINION, use_cache=False)
    except Exception as e:
        return jsonify({"ok": False, "error": f"Bedrock error (auditor opinion): {e}"}), 500

    auditor_opinion = opinion_result.get("auditor_opinion", {})
    logger.info("[Extract Auditor Opinion] Complete: report_type=%s, qualified=%s", 
                auditor_opinion.get("report_type", "unknown"),
                auditor_opinion.get("qualified_opinion", "unknown"))

    # ===========================================
    # Phase 2: Extract vendor controls (batched, from tables)
    # ===========================================
    batch_size = app.config["BATCH_SIZE"]

    merged_controls: List[Dict[str, Any]] = []
    cursor = ""
    passes = 0

    while True:
        instruction = INSTRUCTION_CONTROLS.format(batch_size=batch_size, cursor=cursor)
        try:
            result = invoke_bedrock_with_retry(system_prompt_tables, instruction)
        except Exception as e:
            return jsonify({"ok": False, "error": f"Bedrock error (controls): {e}"}), 500

        controls = (result.get("extraction", {}) or {}).get("controls", [])
        merged_controls = merge_controls(merged_controls, controls)

        meta = result.get("meta", {}) or {}
        cursor = meta.get("last_control_id") or meta.get("last_index") or ""
        has_more = bool(meta.get("has_more"))

        passes += 1
        logger.info("[Extract Vendor Controls] Pass %d: got %d controls, has_more=%s", passes, len(controls), has_more)

        if not has_more:
            break
        time.sleep(0.2)

    logger.info("[Extract Vendor Controls] Complete: %d total controls in %d passes", len(merged_controls), passes)

    # ===========================================
    # Phase 3: Extract exceptions (from tables)
    # ===========================================
    logger.info("[Extract Exceptions] Starting...")
    
    try:
        exceptions_result = invoke_bedrock_with_retry(system_prompt_tables, INSTRUCTION_EXCEPTIONS)
    except Exception as e:
        return jsonify({"ok": False, "error": f"Bedrock error (exceptions): {e}"}), 500

    exceptions = exceptions_result.get("exceptions", [])
    logger.info("[Extract Exceptions] Found %d exceptions", len(exceptions))

    # ===========================================
    # Phase 4: Extract subservice organization controls (batched, from tables)
    # ===========================================
    logger.info("[Extract Subservice] Starting subservice controls extraction (batched)...")
    
    merged_subservice: List[Dict[str, Any]] = []
    subservice_cursor = ""
    subservice_passes = 0

    while True:
        instruction = INSTRUCTION_SUBSERVICE.format(batch_size=batch_size, cursor=subservice_cursor)
        try:
            subservice_result = invoke_bedrock_with_retry(system_prompt_tables, instruction)
        except Exception as e:
            return jsonify({"ok": False, "error": f"Bedrock error (subservice): {e}"}), 500

        controls = subservice_result.get("subservice_controls", [])
        merged_subservice.extend(controls)

        meta = subservice_result.get("meta", {}) or {}
        subservice_cursor = meta.get("last_control_id") or ""
        has_more = bool(meta.get("has_more"))

        subservice_passes += 1
        logger.info("[Extract Subservice] Pass %d: got %d controls, has_more=%s", subservice_passes, len(controls), has_more)

        if not has_more:
            break
        time.sleep(0.2)

    subservice_controls = merged_subservice
    logger.info("[Extract Subservice] Complete: %d total controls in %d passes", len(subservice_controls), subservice_passes)

    # ===========================================
    # Phase 5: Extract user entity controls (batched, from tables)
    # ===========================================
    logger.info("[Extract User Entity] Starting user entity controls extraction (batched)...")
    
    merged_user_entity: List[Dict[str, Any]] = []
    user_entity_cursor = ""
    user_entity_passes = 0

    while True:
        instruction = INSTRUCTION_USER_ENTITY.format(batch_size=batch_size, cursor=user_entity_cursor)
        try:
            user_entity_result = invoke_bedrock_with_retry(system_prompt_tables, instruction)
        except Exception as e:
            return jsonify({"ok": False, "error": f"Bedrock error (user entity): {e}"}), 500

        controls = user_entity_result.get("user_entity_controls", [])
        merged_user_entity.extend(controls)

        meta = user_entity_result.get("meta", {}) or {}
        user_entity_cursor = meta.get("last_description") or ""
        has_more = bool(meta.get("has_more"))

        user_entity_passes += 1
        logger.info("[Extract User Entity] Pass %d: got %d controls, has_more=%s", user_entity_passes, len(controls), has_more)

        if not has_more:
            break
        time.sleep(0.2)

    user_entity_controls = merged_user_entity
    logger.info("[Extract User Entity] Complete: %d total controls in %d passes", len(user_entity_controls), user_entity_passes)

    # ===========================================
    # Phase 6: Extract criteria mappings (from tables)
    # ===========================================
    logger.info("[Extract Criteria] Starting criteria extraction...")
    
    try:
        criteria_result = invoke_bedrock_with_retry(system_prompt_tables, INSTRUCTION_CRITERIA)
    except Exception as e:
        return jsonify({"ok": False, "error": f"Bedrock error (criteria): {e}"}), 500

    criteria_mappings = criteria_result.get("criteria_mappings", [])
    logger.info("[Extract Criteria] Found %d criteria", len(criteria_mappings))

    # ===========================================
    # Phase 7: Merge criteria into vendor controls
    # ===========================================
    final_controls = merge_criteria_into_controls(merged_controls, criteria_mappings)
    logger.info("[Merge] Merged criteria into %d controls", len(final_controls))

    # ===========================================
    # Phase 8: Build DOCX with all content
    # ===========================================
    rid = str(uuid.uuid4())[:8]
    docx_path = _report_path_docx(rid)
    try:
        build_report_docx(
            auditor_opinion=auditor_opinion,
            subservice_controls=subservice_controls,
            user_entity_controls=user_entity_controls,
            vendor_controls=final_controls,
            exceptions=exceptions,
            criteria_mappings=criteria_mappings,
            out_path=docx_path
        )
    except Exception as e:
        logger.exception("Failed to build DOCX: %s", e)

    docx_ok = os.path.exists(docx_path)
    elapsed = round(time.time() - t0, 2)

    return jsonify({
        "ok": True,
        "result": {
            "extraction": {
                "auditor_opinion": auditor_opinion,
                "vendor_controls": final_controls,
                "exceptions": exceptions,
                "subservice_controls": subservice_controls,
                "user_entity_controls": user_entity_controls,
            },
            "criteria_mappings": criteria_mappings,
            "meta": {
                "vendor_controls_found": len(final_controls),
                "exceptions_found": len(exceptions),
                "subservice_controls_found": len(subservice_controls),
                "user_entity_controls_found": len(user_entity_controls),
                "criteria_found": len(criteria_mappings),
                "vendor_passes": passes,
                "subservice_passes": subservice_passes,
                "user_entity_passes": user_entity_passes,
            }
        },
        "report_id": rid if docx_ok else None,
        "elapsed_sec": elapsed,
        "passes": passes,
    })


@app.get("/download/docx/<rid>")
def download_docx(rid: str):
    docx_path = _report_path_docx(rid)
    if not os.path.exists(docx_path):
        return jsonify({"ok": False, "error": "DOCX not found"}), 404
    return send_file(docx_path, as_attachment=True, download_name=f"soc2_report_{rid}.docx")


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", "5000")), debug=True)
